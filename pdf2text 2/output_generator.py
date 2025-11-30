"""
Output Generation Module
Export optimized content to Word (.docx) and Markdown (.md) formats
Preserves structure including headings, lists, tables, and formatting
"""
from pathlib import Path
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import pandas as pd
from loguru import logger
import json
import csv


class WordGenerator:
    """Generate Word documents from processed content"""
    
    def __init__(self):
        """Initialize Word document generator"""
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup document styles"""
        try:
            # Set default font
            style = self.doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Setup heading styles
            for i in range(1, 4):
                heading_style = self.doc.styles[f'Heading {i}']
                heading_style.font.name = 'Calibri'
                heading_style.font.bold = True
                heading_style.font.color.rgb = RGBColor(0, 0, 0)
                
        except Exception as e:
            logger.warning(f"Error setting up styles: {e}")
    
    def add_title(self, title: str):
        """Add document title"""
        heading = self.doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_heading(self, text: str, level: int = 1):
        """Add heading"""
        self.doc.add_heading(text, level=level)
    
    def add_paragraph(self, text: str, style: Optional[str] = None):
        """Add paragraph"""
        if text:
            para = self.doc.add_paragraph(text, style=style)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def add_table(self, table_data: List[List[str]], has_header: bool = True):
        """
        Add table to document
        
        Args:
            table_data: 2D list of table data
            has_header: Whether first row is header
        """
        if not table_data:
            return
        
        try:
            rows = len(table_data)
            cols = len(table_data[0]) if table_data else 0
            
            if rows == 0 or cols == 0:
                return
            
            # Create table
            table = self.doc.add_table(rows=rows, cols=cols)
            table.style = 'Light Grid Accent 1'
            
            # Populate table
            for i, row_data in enumerate(table_data):
                row = table.rows[i]
                for j, cell_data in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = str(cell_data) if cell_data else ''
                    
                    # Format header row
                    if i == 0 and has_header:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
            
            # Add spacing after table
            self.doc.add_paragraph()
            
        except Exception as e:
            logger.error(f"Error adding table: {e}")
    
    def add_metadata(self, metadata: Dict):
        """Add metadata section"""
        self.add_heading("Document Information", level=2)
        
        for key, value in metadata.items():
            if value and str(value) != 'Unknown':
                self.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
        
        self.doc.add_page_break()
    
    def add_page_break(self):
        """Add page break"""
        self.doc.add_page_break()
    
    def generate_from_pdf_result(
        self,
        pdf_result: Dict,
        include_metadata: bool = True,
        include_tables: bool = True,
        title: Optional[str] = None
    ):
        """
        Generate Word document from PDF parsing result
        
        Args:
            pdf_result: Dictionary from PDF parser
            include_metadata: Whether to include metadata
            include_tables: Whether to include tables
            title: Document title
        """
        try:
            # Add title
            doc_title = title or pdf_result.get('metadata', {}).get('title', 'Document')
            self.add_title(doc_title)
            
            # Add metadata
            if include_metadata and pdf_result.get('metadata'):
                self.add_metadata(pdf_result['metadata'])
            
            # Add main content
            self.add_heading("Content", level=1)
            
            # Process pages
            for page in pdf_result.get('pages', []):
                page_num = page.get('page_number', 0)
                
                # Add page heading
                self.add_heading(f"Page {page_num}", level=2)
                
                # Add paragraphs
                for para in page.get('paragraphs', []):
                    if para and para.strip():
                        self.add_paragraph(para)
                
                # Add tables
                if include_tables:
                    for table in page.get('tables', []):
                        table_num = table.get('table_number', 0)
                        self.add_heading(f"Table {table_num}", level=3)
                        self.add_table(table.get('data', []))
            
            logger.info("Word document generated successfully")
            
        except Exception as e:
            logger.error(f"Error generating Word document: {e}")
            raise
    
    def save(self, output_path: str):
        """Save document to file"""
        try:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            self.doc.save(str(output_path))
            logger.info(f"Word document saved to: {output_path}")
            
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error saving Word document: {e}")
            raise


class MarkdownGenerator:
    """Generate Markdown documents from processed content"""
    
    def __init__(self):
        """Initialize Markdown generator"""
        self.content = []
    
    def add_title(self, title: str):
        """Add document title"""
        self.content.append(f"# {title}\n")
    
    def add_heading(self, text: str, level: int = 2):
        """Add heading"""
        prefix = '#' * level
        self.content.append(f"{prefix} {text}\n")
    
    def add_paragraph(self, text: str):
        """Add paragraph"""
        if text:
            self.content.append(f"{text}\n")
    
    def add_table(self, table_data: List[List[str]], has_header: bool = True):
        """
        Add table in Markdown format
        
        Args:
            table_data: 2D list of table data
            has_header: Whether first row is header
        """
        if not table_data:
            return
        
        try:
            # Add header row
            if has_header and len(table_data) > 0:
                header = table_data[0]
                self.content.append('| ' + ' | '.join(str(cell) for cell in header) + ' |')
                self.content.append('|' + '|'.join(['---' for _ in header]) + '|')
                
                # Add data rows
                for row in table_data[1:]:
                    self.content.append('| ' + ' | '.join(str(cell) for cell in row) + ' |')
            else:
                # No header
                for row in table_data:
                    self.content.append('| ' + ' | '.join(str(cell) for cell in row) + ' |')
            
            self.content.append('')  # Empty line after table
            
        except Exception as e:
            logger.error(f"Error adding table: {e}")
    
    def add_metadata(self, metadata: Dict):
        """Add metadata section"""
        self.add_heading("Document Information", level=2)
        
        for key, value in metadata.items():
            if value and str(value) != 'Unknown':
                self.add_paragraph(f"**{key.replace('_', ' ').title()}**: {value}")
        
        self.add_paragraph("---\n")
    
    def add_horizontal_rule(self):
        """Add horizontal rule"""
        self.content.append("---\n")
    
    def add_code_block(self, code: str, language: str = ''):
        """Add code block"""
        self.content.append(f"```{language}")
        self.content.append(code)
        self.content.append("```\n")
    
    def generate_from_pdf_result(
        self,
        pdf_result: Dict,
        include_metadata: bool = True,
        include_tables: bool = True,
        title: Optional[str] = None
    ):
        """
        Generate Markdown document from PDF parsing result
        
        Args:
            pdf_result: Dictionary from PDF parser
            include_metadata: Whether to include metadata
            include_tables: Whether to include tables
            title: Document title
        """
        try:
            # Add title
            doc_title = title or pdf_result.get('metadata', {}).get('title', 'Document')
            self.add_title(doc_title)
            
            # Add metadata
            if include_metadata and pdf_result.get('metadata'):
                self.add_metadata(pdf_result['metadata'])
            
            # Add main content
            self.add_heading("Content", level=2)
            
            # Process pages
            for page in pdf_result.get('pages', []):
                page_num = page.get('page_number', 0)
                
                # Add page heading
                self.add_heading(f"Page {page_num}", level=3)
                
                # Add paragraphs
                for para in page.get('paragraphs', []):
                    if para and para.strip():
                        self.add_paragraph(para)
                        self.add_paragraph('')  # Empty line between paragraphs
                
                # Add tables
                if include_tables:
                    for table in page.get('tables', []):
                        table_num = table.get('table_number', 0)
                        self.add_heading(f"Table {table_num}", level=4)
                        self.add_table(table.get('data', []))
            
            logger.info("Markdown document generated successfully")
            
        except Exception as e:
            logger.error(f"Error generating Markdown document: {e}")
            raise
    
    def get_content(self) -> str:
        """Get generated Markdown content"""
        return '\n'.join(self.content)
    
    def save(self, output_path: str):
        """Save Markdown to file"""
        try:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(self.get_content())
            
            logger.info(f"Markdown document saved to: {output_path}")
            
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error saving Markdown document: {e}")
            raise


class DataExporter:
    """Export table data to CSV and JSON formats"""
    
    @staticmethod
    def export_tables_to_csv(tables: List[Dict], output_dir: str) -> List[str]:
        """
        Export tables to CSV files
        
        Args:
            tables: List of table dictionaries
            output_dir: Output directory
            
        Returns:
            List of output file paths
        """
        output_paths = []
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        try:
            for idx, table in enumerate(tables, 1):
                table_data = table.get('data', [])
                if not table_data:
                    continue
                
                output_path = output_dir / f"table_{idx}.csv"
                
                with open(output_path, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    writer.writerows(table_data)
                
                output_paths.append(str(output_path))
                logger.info(f"Table {idx} exported to CSV: {output_path}")
            
            return output_paths
            
        except Exception as e:
            logger.error(f"Error exporting tables to CSV: {e}")
            return output_paths
    
    @staticmethod
    def export_tables_to_json(tables: List[Dict], output_path: str) -> str:
        """
        Export all tables to a single JSON file
        
        Args:
            tables: List of table dictionaries
            output_path: Output file path
            
        Returns:
            Output file path
        """
        try:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(tables, f, ensure_ascii=False, indent=2)
            
            logger.info(f"Tables exported to JSON: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error exporting tables to JSON: {e}")
            raise
    
    @staticmethod
    def export_text(text: str, output_path: str) -> str:
        """
        Export text to file
        
        Args:
            text: Text content
            output_path: Output file path
            
        Returns:
            Output file path
        """
        try:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            
            logger.info(f"Text exported to: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error exporting text: {e}")
            raise
